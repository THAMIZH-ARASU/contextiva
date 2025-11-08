"""Text extraction service for multiple file formats."""
import asyncio
from pathlib import Path
from typing import BinaryIO

from bs4 import BeautifulSoup
from docx import Document
from PyPDF2 import PdfReader

from src.shared.utils.errors import TextExtractionError


class TextExtractor:
    """Service for extracting text from various file formats."""

    async def extract(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from file based on file extension.

        Args:
            file_content: File content as bytes
            filename: Name of the file with extension

        Returns:
            Extracted text content

        Raises:
            TextExtractionError: If extraction fails or format is unsupported
        """
        file_extension = Path(filename).suffix.lower()

        try:
            if file_extension == ".md":
                return await self.extract_markdown(file_content)
            elif file_extension == ".pdf":
                return await self.extract_pdf(file_content)
            elif file_extension == ".docx":
                return await self.extract_docx(file_content)
            elif file_extension == ".html":
                return await self.extract_html(file_content)
            else:
                raise TextExtractionError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            raise TextExtractionError(f"Failed to extract text from {filename}: {str(e)}")

    async def extract_markdown(self, file_content: bytes) -> str:
        """
        Extract text from Markdown file.

        Args:
            file_content: Markdown file content as bytes

        Returns:
            Raw text content
        """
        return file_content.decode("utf-8")

    async def extract_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text from all pages

        Raises:
            TextExtractionError: If PDF is corrupted or cannot be read
        """
        try:
            # Run CPU-intensive PDF parsing in thread pool
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, self._extract_pdf_sync, file_content)
            return text
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_pdf_sync(self, file_content: bytes) -> str:
        """Synchronous PDF extraction (run in thread pool)."""
        from io import BytesIO

        pdf_file = BytesIO(file_content)
        reader = PdfReader(pdf_file)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    async def extract_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text from paragraphs and tables

        Raises:
            TextExtractionError: If DOCX is corrupted or cannot be read
        """
        try:
            # Run CPU-intensive DOCX parsing in thread pool
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, self._extract_docx_sync, file_content)
            return text
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")

    def _extract_docx_sync(self, file_content: bytes) -> str:
        """Synchronous DOCX extraction (run in thread pool)."""
        from io import BytesIO

        docx_file = BytesIO(file_content)
        doc = Document(docx_file)
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    async def extract_html(self, file_content: bytes) -> str:
        """
        Extract text from HTML file.

        Args:
            file_content: HTML file content as bytes

        Returns:
            Extracted text with tags removed

        Raises:
            TextExtractionError: If HTML is malformed or cannot be parsed
        """
        try:
            html_content = file_content.decode("utf-8")
            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            return text
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from HTML: {str(e)}")
